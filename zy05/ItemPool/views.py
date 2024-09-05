from django.http import HttpResponse
from .models import *   #导入模型
import random
import os
def makePaperConent(request):
    #从请求中获得使用的试卷模板的ID
    pid=int(request.GET['papertemplateid'])
    #获得ID对应的试卷模板对象，从其中获取试卷设置用于抽取试题
    pt=paperTemplate.objects.get(id=pid)
    #随机抽取试题的基本思路：
    #首从数据库按题型返回试题ID，返回的对象为QuerySet，
    #根据QuerySet的count()值，调用random.randint()获得该范围内的随机整数。
    #用随机整数作为索引获得试题ID值，将其加入试题列表，加入时要避免重复  
    item_ids=getItems(pt.typeOneCount,'单项选择题')  #抽取单项选择题 
    item_ids+=getItems(pt.typeTwoCount,'基本操作题')  #抽取基本操作题
    item_ids+=getItems(pt.typeThreeCount,'简单应用题')  #抽取简单应用题
    item_ids+=getItems(pt.typeFourCount,'综合应用题')  #抽取综合应用题    
    return HttpResponse('%s' % item_ids)

def getItems(count,key):#根据数量和题型抽取试题
    td=[]#初始的空白试题列表，保存已抽出的试题ID
    n=0 #用于对已抽出的试题ID进行计数
    #按题型获得试题ID的QuerySet
    ts=testItem.objects.filter(type__name=key).values('id')
    while n<count:    #抽取试题
        x=random.randint(0,ts.count()-1)
        while ts[x]['id'] in td:
            #如果随机抽取的ID与已选择ID重复，则重新生成随机数
            x=random.randint(0,ts.count()-1)
        td.append(ts[x]['id'])#将不重复的ID加入列表
        n=n+1
    return td

from django.template.response import TemplateResponse
from docx import Document
from docx.shared import Cm

def getPaperIndex(request): #返回试卷导出页面
    pcs=paperContent.objects.order_by('name')#获取试卷数据
    return TemplateResponse(request,'showpaper.html',{'pcs':pcs})

def getPaper(request):#生成试卷预览内容和Word文件
    pid=int(request.GET['paperid'])         #获得试卷ID
    pc=paperContent.objects.get(id=pid)     #获得ID对应的试卷对象
    #itemList保存试卷试题ID列表，后面会通过ID从模型获取试题
    #在抽取试题时，按照单项选择题、基本操作题、简单应用题和综合应用题的顺序，
    #后面会按该顺序从itemList获得试题的ID来输出试卷详细内容。
    itemList=eval(pc.content)
    #paper变量保存试卷预览内容
    ptitle="%s_%s"%(pc.template.name,pc.name)
    paper="<h1>%s</h1>"%ptitle
    #document变量用于将试卷内容写入Word文档
    document = Document()   #新建空白Word文档    
    document.add_heading(ptitle,level=1)#作为一级标题写入Word文档
    #变量path保存图片文件的系统路径，将图片写入Word文档时用该路径获取图片
    path=os.path.dirname(os.path.abspath(__file__))+'\\static\\'
    #PT为试卷模板对象，从其中获得各类型试题的数量和分值
    pt=paperTemplate.objects.get(id=pc.template.id)  
    
    #输出单项选择题
    t='一、输出单项选择题。（共%s小题，每小题%s分）'%(pt.typeOneCount,pt.typeOneScore)
    paper+='<p>%s</p>'%t        #添加预览内容
    document.add_paragraph(t)   #向Word文件添加段落
    index=1
    while index<=pt.typeOneCount:
        #根据试题ID，通过模型对象获得试题
        item=testItem.objects.get(id=itemList[index-1])
        t='%s.%s'%(index,item.question)
        paper+='%s<br>'%t
        document.add_paragraph(t)
        if item.picture:
            #有图片时，注意预览内容使用图片的URL，
            #写入Word文档时，使用图片的本地系统路径
            src='/static/%s'%item.picture
            paper+='<img src="%s"><br>'%src
            document.add_picture(path+'%s'%item.picture, width=Cm(2))
        op=eval(item.options)#获得单选题的选项字典
        paper+='（A）%s<br>（B）%s<br>（C）%s<br>（D）%s<br>'%\
            (op['A'],op['B'],op['C'],op['D'])
        t='（A）%s（B）%s（C）%s（D）%s'%(op['A'],op['B'],op['C'],op['D'])
        document.add_paragraph(t)
        index=index+1

    #输出基本操作题
    t='二、基本操作题。（共%s小题，每小题%s分）'%(pt.typeTwoCount,pt.typeTwoScore)
    paper+='<p>%s</p>'%t
    document.add_paragraph(t)
    n=index
    index=1
    while index<=pt.typeTwoCount:
        item=testItem.objects.get(id=itemList[n-1])
        t='%s.%s'%(index,item.question)
        paper+='%s<br>'%t
        document.add_paragraph(t)
        if item.picture:
            src='/static/%s'%item.picture
            paper+='<img src="%s"><br>'%src
            document.add_picture(path+'%s'%item.picture, width=Cm(2))
        n=n+1
        index=index+1

    #输出简单应用题
    t='三、简单应用题。（共%s小题，每小题%s分）'% (pt.typeThreeCount,pt.typeThreeScore)
    paper+='<p>%s</p>'%t
    index=1
    while index<=pt.typeThreeCount:
        item=testItem.objects.get(id=itemList[n-1])
        t='%s.%s'%(index,item.question)
        paper+='%s<br>'%t
        document.add_paragraph(t)
        if item.picture:
            src='/static/%s'%item.picture
            paper+='<img src="%s"><br>'%src
            document.add_picture(path+'%s'%item.picture, width=Cm(2))
        n=n+1
        index=index+1  
        
    #输出综合应用题
    t='四、综合应用题。（共%s小题，每小题%s分）'%(pt.typeFourCount,pt.typeFourScore)
    paper+='<p>%s</p>'%t
    index=1
    while index<=pt.typeFourCount:
        item=testItem.objects.get(id=itemList[n-1])
        t='%s.%s'%(index,item.question)
        paper+='%s<br>'%t
        document.add_paragraph(t)
        if item.picture:
            src='/static/%s'%item.picture
            paper+='<img src="%s"><br>'%src
            document.add_picture(path+'%s'%item.picture, width=Cm(2))
        n=n+1
        index=index+1

    document.add_page_break()#结束Word文档
    filename='%s_%s.docx'%(pc.template.name,pc.name)
    document.save(path+filename)#Word文档写入系统文件
    request.session['paperfile']=filename#在会话中保存试卷名称，用于生成下载地址
    return HttpResponse(paper)

def getDoc(request):#向客户端返回试卷Word文件下载地址
    filename=request.session['paperfile']
    url='<a href="/static/%s">%s</a>' %(filename,filename)#生成下载地址
    return HttpResponse(url)